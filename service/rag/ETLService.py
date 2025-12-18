"""
RAG 数据清洗服务 (ETL Service)
这是 RAG 的"地基"，数据质量直接决定系统上限。

功能包括：
1. 格式统一：处理 PDF、Word、Markdown、HTML 等不同格式
2. PDF 特殊处理：解决表格、多栏布局解析错乱的问题
3. 去噪：去除页眉、页脚、HTML 标签、无意义的特殊字符
"""
import os
import re
from typing import List, Optional, Dict, Any
from pathlib import Path


class ETLService:
    """数据清洗服务类"""

    def __init__(self, use_unstructured: bool = False, unstructured_api_key: Optional[str] = None):
        """
        初始化 ETL 服务
        
        Args:
            use_unstructured: 是否使用 Unstructured 库处理 PDF（需要额外安装）
            unstructured_api_key: Unstructured API 密钥（如果使用云服务）
        """
        self.use_unstructured = use_unstructured
        self.unstructured_api_key = unstructured_api_key
        self._init_loaders()

    def _init_loaders(self):
        """初始化各种文档加载器"""
        # PDF 处理
        try:
            import pdfplumber
            self.pdfplumber_available = True
        except ImportError:
            self.pdfplumber_available = False
            print("警告: pdfplumber 未安装，PDF 处理功能受限。建议安装: pip install pdfplumber")

        try:
            import PyPDF2
            self.pypdf2_available = True
        except ImportError:
            self.pypdf2_available = False

        # Word 处理 (.docx)
        try:
            from docx import Document
            self.docx_available = True
        except ImportError:
            self.docx_available = False
            print("警告: python-docx 未安装，Word 处理功能不可用。建议安装: pip install python-docx")

        # Word 处理 (.doc - 旧版格式)
        try:
            import textract
            self.textract_available = True
        except ImportError:
            self.textract_available = False
            # textract 不是必需的，因为大多数文件都是 .docx 格式

        # 备用方案：docx2txt（可能对某些 .doc 文件有效）
        try:
            import docx2txt
            self.docx2txt_available = True
        except ImportError:
            self.docx2txt_available = False

        # 备用方案：使用 antiword 或 catdoc（需要系统安装）
        # 这里我们主要依赖 python-docx，对于 .doc 文件给出更好的错误提示

        # HTML 处理
        try:
            from bs4 import BeautifulSoup
            self.bs4_available = True
        except ImportError:
            self.bs4_available = False
            print("警告: BeautifulSoup4 未安装，HTML 处理功能受限。建议安装: pip install beautifulsoup4")

        # Markdown 处理（Python 内置支持，但可以使用 markdown 库增强）
        try:
            import markdown
            self.markdown_available = True
        except ImportError:
            self.markdown_available = False

        # Unstructured 处理（高级 PDF 处理）
        if self.use_unstructured:
            try:
                from unstructured.partition.auto import partition
                self.unstructured_available = True
            except ImportError:
                self.unstructured_available = False
                print("警告: unstructured 未安装，高级 PDF 处理功能不可用。建议安装: pip install unstructured")

    def load_document(self, file_path: str) -> Dict[str, Any]:
        """
        加载文档，自动识别格式并提取内容
        
        Args:
            file_path: 文档路径
            
        Returns:
            包含文档内容和元数据的字典
            {
                'content': str,  # 提取的文本内容
                'format': str,    # 文档格式 (pdf, docx, html, md, txt)
                'metadata': dict  # 文档元数据
            }
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        suffix = file_path.suffix.lower()

        # 检测实际文件格式（可能扩展名不匹配）
        actual_format = self._detect_file_format(file_path, suffix)

        if suffix == '.pdf' or actual_format == 'pdf':
            return self._load_pdf(file_path)
        elif suffix == '.docx' or (actual_format == 'docx' and suffix == '.docx'):
            return self._load_word(file_path)
        elif suffix == '.doc' or (actual_format == 'doc' and suffix == '.doc'):
            return self._load_word_old(file_path)
        elif suffix in ['.html', '.htm']:
            return self._load_html(file_path)
        elif suffix == '.md':
            return self._load_markdown(file_path)
        elif suffix == '.txt':
            return self._load_text(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")

    def _load_pdf(self, file_path: Path) -> Dict[str, Any]:
        """加载 PDF 文件"""
        content_parts = []
        metadata = {'pages': 0, 'method': 'unknown'}

        # 优先使用 Unstructured（如果可用且启用）
        if self.use_unstructured and self.unstructured_available:
            try:
                from unstructured.partition.auto import partition
                elements = partition(str(file_path))
                content_parts = [str(elem) for elem in elements]
                content = '\n\n'.join(content_parts)
                metadata['method'] = 'unstructured'
                metadata['pages'] = len([e for e in elements if hasattr(e, 'metadata')])
                return {
                    'content': content,
                    'format': 'pdf',
                    'metadata': metadata
                }
            except Exception as e:
                print(f"Unstructured 处理失败，尝试其他方法: {e}")

        # 使用 pdfplumber（推荐，对表格支持好）
        if self.pdfplumber_available:
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    pages_content = []
                    for page in pdf.pages:
                        # 提取文本
                        text = page.extract_text()
                        if text:
                            pages_content.append(text)

                        # 提取表格
                        tables = page.extract_tables()
                        for table in tables:
                            if table:
                                table_text = self._table_to_text(table)
                                pages_content.append(table_text)

                    content = '\n\n'.join(pages_content)
                    metadata['method'] = 'pdfplumber'
                    metadata['pages'] = len(pdf.pages)
                    return {
                        'content': content,
                        'format': 'pdf',
                        'metadata': metadata
                    }
            except Exception as e:
                print(f"pdfplumber 处理失败: {e}")

        # 使用 PyPDF2（备用方案）
        if self.pypdf2_available:
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    pages_content = []
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            pages_content.append(text)
                    content = '\n\n'.join(pages_content)
                    metadata['method'] = 'pypdf2'
                    metadata['pages'] = len(pdf_reader.pages)
                    return {
                        'content': content,
                        'format': 'pdf',
                        'metadata': metadata
                    }
            except Exception as e:
                print(f"PyPDF2 处理失败: {e}")

        raise RuntimeError("无法处理 PDF 文件，请安装 pdfplumber 或 PyPDF2")

    def _table_to_text(self, table: List[List[str]]) -> str:
        """将表格转换为文本格式"""
        if not table:
            return ""

        rows = []
        for row in table:
            # 过滤 None 值并转换为字符串
            clean_row = [str(cell) if cell else "" for cell in row]
            rows.append(" | ".join(clean_row))

        return "\n".join(rows)

    def _detect_file_format(self, file_path: Path, suffix: str) -> str:
        """
        检测文件的实际格式（可能扩展名不匹配）
        
        Args:
            file_path: 文件路径
            suffix: 文件扩展名
            
        Returns:
            检测到的文件格式
        """
        import zipfile

        # 检查是否为 ZIP 格式（.docx 实际上是 ZIP）
        if suffix in ['.docx', '.doc']:
            try:
                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    if 'word/document.xml' in zip_ref.namelist():
                        return 'docx'
            except (zipfile.BadZipFile, Exception):
                # 不是 ZIP 格式，可能是旧版 .doc
                if suffix == '.doc':
                    return 'doc'
                # 如果扩展名是 .docx 但实际不是 ZIP，可能是旧版 .doc
                return 'doc'

        return suffix.lstrip('.')

    def _load_word(self, file_path: Path) -> Dict[str, Any]:
        """加载 Word 文件 (.docx 格式)"""
        if not self.docx_available:
            raise ImportError("请安装 python-docx: pip install python-docx")

        try:
            from docx import Document
            # 验证文件是否为有效的 .docx 文件
            # .docx 文件实际上是一个 ZIP 压缩包
            import zipfile
            try:
                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    # 检查是否包含必要的文件
                    if 'word/document.xml' not in zip_ref.namelist():
                        raise ValueError("文件不是有效的 .docx 格式")
            except zipfile.BadZipFile:
                # 可能是旧版 .doc 格式，尝试使用 .doc 处理方法
                print(f"警告: 文件 '{file_path.name}' 不是有效的 .docx 格式，尝试作为 .doc 格式处理...")
                try:
                    return self._load_word_old(file_path)
                except Exception as e:
                    raise ValueError(
                        f"文件 '{file_path.name}' 不是有效的 .docx 格式，也无法作为 .doc 格式处理。"
                        f"错误详情: {str(e)}。"
                        f"请将文件转换为 .docx 格式，或使用支持 .doc 格式的工具（如 textract）。"
                    )

            doc = Document(file_path)

            content_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)

            # 提取表格
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(cells))
                if table_text:
                    content_parts.append("\n".join(table_text))

            content = '\n\n'.join(content_parts)

            return {
                'content': content,
                'format': 'docx',
                'metadata': {'paragraphs': len(doc.paragraphs), 'tables': len(doc.tables)}
            }
        except zipfile.BadZipFile as e:
            # 尝试作为 .doc 格式处理
            print(f"警告: 文件 '{file_path.name}' 不是有效的 .docx 格式，尝试作为 .doc 格式处理...")
            try:
                return self._load_word_old(file_path)
            except Exception as e2:
                raise ValueError(
                    f"无法读取文件 '{file_path.name}'：文件不是有效的 .docx 格式，也无法作为 .doc 格式处理。"
                    f"错误详情: {str(e2)}。"
                    f"请将文件转换为 .docx 格式，或使用支持 .doc 格式的工具（如 textract）。"
                )
        except Exception as e:
            # 检查是否是 zipfile 相关错误
            if "not a zip file" in str(e).lower() or "badzipfile" in str(e).lower():
                print(f"警告: 文件 '{file_path.name}' 不是有效的 .docx 格式，尝试作为 .doc 格式处理...")
                try:
                    return self._load_word_old(file_path)
                except Exception as e2:
                    raise ValueError(
                        f"无法读取文件 '{file_path.name}'：文件不是有效的 .docx 格式，也无法作为 .doc 格式处理。"
                        f"错误详情: {str(e2)}。"
                        f"请将文件转换为 .docx 格式，或使用支持 .doc 格式的工具（如 textract）。"
                    )
            raise RuntimeError(f"读取 Word 文件时出错: {str(e)}")

    def _load_word_old(self, file_path: Path) -> Dict[str, Any]:
        """加载旧版 Word 文件 (.doc 格式)"""
        # 尝试使用 textract
        if self.textract_available:
            try:
                import textract
                content = textract.process(str(file_path)).decode('utf-8')
                return {
                    'content': content,
                    'format': 'doc',
                    'metadata': {'method': 'textract'}
                }
            except Exception as e:
                print(f"textract 处理失败: {e}")
                # 继续尝试其他方法

        # 尝试使用 docx2txt（如果可用）
        try:
            import docx2txt
            content = docx2txt.process(str(file_path))
            if content:
                return {
                    'content': content,
                    'format': 'doc',
                    'metadata': {'method': 'docx2txt'}
                }
        except ImportError:
            pass
        except Exception as e:
            print(f"docx2txt 处理失败: {e}")

        # 如果所有方法都不可用，提供清晰的错误信息
        error_msg = (
            f"无法处理旧版 .doc 格式文件 '{file_path.name}'。\n"
            f"请选择以下方案之一：\n"
            f"1. 将文件转换为 .docx 格式（推荐）：在 Word 中打开文件，选择 '另存为' -> 选择 .docx 格式\n"
            f"2. 安装 textract: pip install textract（需要系统依赖，Windows 上可能需要额外配置）\n"
            f"3. 安装 docx2txt: pip install docx2txt（可能对某些 .doc 文件有效）\n"
            f"4. 使用在线工具或其他软件手动转换文件格式"
        )
        raise ImportError(error_msg)

    def _load_html(self, file_path: Path) -> Dict[str, Any]:
        """加载 HTML 文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()

        # 使用 BeautifulSoup 解析（如果可用）
        if self.bs4_available:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')

            # 移除脚本和样式标签
            for script in soup(["script", "style", "meta", "link"]):
                script.decompose()

            # 提取文本
            text = soup.get_text(separator='\n', strip=True)
        else:
            # 简单的 HTML 标签移除（备用方案）
            text = self._remove_html_tags(html_content)

        return {
            'content': text,
            'format': 'html',
            'metadata': {}
        }

    def _load_markdown(self, file_path: Path) -> Dict[str, Any]:
        """加载 Markdown 文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        return {
            'content': content,
            'format': 'md',
            'metadata': {}
        }

    def _load_text(self, file_path: Path) -> Dict[str, Any]:
        """加载纯文本文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        return {
            'content': content,
            'format': 'txt',
            'metadata': {}
        }

    def clean_text(
            self,
            text: str,
            remove_headers_footers: bool = True,
            remove_html_tags: bool = True,
            remove_special_chars: bool = True,
            remove_extra_whitespace: bool = True,
            custom_patterns: Optional[List[str]] = None
    ) -> str:
        """
        文本去噪处理
        
        Args:
            text: 原始文本
            remove_headers_footers: 是否去除页眉页脚
            remove_html_tags: 是否去除 HTML 标签
            remove_special_chars: 是否去除无意义的特殊字符
            remove_extra_whitespace: 是否去除多余空白
            custom_patterns: 自定义正则表达式模式列表，用于去除特定内容
            
        Returns:
            清洗后的文本
        """
        cleaned_text = text

        # 去除 HTML 标签
        if remove_html_tags:
            cleaned_text = self._remove_html_tags(cleaned_text)

        # 去除页眉页脚（常见模式）
        if remove_headers_footers:
            cleaned_text = self._remove_headers_footers(cleaned_text)

        # 去除无意义的特殊字符
        if remove_special_chars:
            cleaned_text = self._remove_special_chars(cleaned_text)

        # 应用自定义模式
        if custom_patterns:
            for pattern in custom_patterns:
                cleaned_text = re.sub(pattern, '', cleaned_text, flags=re.IGNORECASE | re.MULTILINE)

        # 去除多余空白
        if remove_extra_whitespace:
            cleaned_text = self._normalize_whitespace(cleaned_text)

        return cleaned_text.strip()

    def _remove_html_tags(self, text: str) -> str:
        """移除 HTML 标签"""
        # 移除 HTML 注释
        text = re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)
        # 移除 HTML 标签
        text = re.sub(r'<[^>]+>', '', text)
        # 解码 HTML 实体
        try:
            import html
            text = html.unescape(text)
        except:
            pass
        return text

    def _remove_headers_footers(self, text: str) -> str:
        """去除页眉页脚"""
        lines = text.split('\n')
        cleaned_lines = []

        for line in lines:
            line = line.strip()
            # 跳过常见的页眉页脚模式
            # 页码模式：纯数字或 "第 X 页"
            if re.match(r'^(\d+|[第]\s*\d+\s*[页])$', line):
                continue
            # 日期模式：YYYY-MM-DD 或类似格式
            if re.match(r'^\d{4}[-/]\d{1,2}[-/]\d{1,2}$', line):
                continue
            # 短行且包含常见页眉页脚关键词
            if len(line) < 30 and any(keyword in line for keyword in ['第', '页', '共', '页眉', '页脚']):
                continue

            cleaned_lines.append(line)

        return '\n'.join(cleaned_lines)

    def _remove_special_chars(self, text: str) -> str:
        """去除无意义的特殊字符"""
        # 保留常见的有意义字符：字母、数字、中文、基本标点、换行、制表符
        # 移除控制字符（除了换行和制表符）
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)

        # 移除零宽字符
        text = re.sub(r'[\u200b-\u200f\ufeff]', '', text)

        # 移除过多的连续特殊字符（保留单个）
        text = re.sub(r'[^\w\s\u4e00-\u9fff]{3,}', '', text)

        return text

    def _normalize_whitespace(self, text: str) -> str:
        """规范化空白字符"""
        # 将多个连续空格替换为单个空格
        text = re.sub(r' +', ' ', text)
        # 将多个连续换行替换为最多两个换行
        text = re.sub(r'\n{3,}', '\n\n', text)
        # 移除行首行尾空白
        lines = [line.rstrip() for line in text.split('\n')]
        text = '\n'.join(lines)
        return text

    def process_document(
            self,
            file_path: str,
            clean_options: Optional[Dict[str, bool]] = None
    ) -> Dict[str, Any]:
        """
        完整的文档处理流程：加载 -> 清洗 -> 返回
        Args:
            file_path: 文档路径
            clean_options: 清洗选项字典，键为 clean_text 方法的参数名
        Returns:
            处理后的文档数据
            {
                'content': str,      # 清洗后的文本内容
                'format': str,        # 文档格式
                'metadata': dict,     # 文档元数据
                'original_length': int,  # 原始文本长度
                'cleaned_length': int    # 清洗后文本长度
            }
        """
        # 加载文档
        doc_data = self.load_document(file_path)
        original_content = doc_data['content']
        original_length = len(original_content)

        # 设置默认清洗选项
        if clean_options is None:
            clean_options = {
                'remove_headers_footers': True,
                'remove_html_tags': True,
                'remove_special_chars': True,
                'remove_extra_whitespace': True
            }

        # 清洗文本
        cleaned_content = self.clean_text(original_content, **clean_options)
        cleaned_length = len(cleaned_content)

        return {
            'content': cleaned_content,
            'format': doc_data['format'],
            'metadata': {
                **doc_data['metadata'],
                'original_length': original_length,
                'cleaned_length': cleaned_length,
                'reduction_ratio': f"{(1 - cleaned_length / original_length) * 100:.2f}%" if original_length > 0 else "0%"
            }
        }

    def batch_process(
            self,
            file_paths: List[str],
            clean_options: Optional[Dict[str, bool]] = None
    ) -> List[Dict[str, Any]]:
        """
        批量处理多个文档
        Args:
            file_paths: 文档路径列表
            clean_options: 清洗选项
        Returns:
            处理后的文档数据列表
        """
        results = []
        for file_path in file_paths:
            try:
                result = self.process_document(file_path, clean_options)
                results.append(result)
            except Exception as e:
                print(f"处理文件 {file_path} 时出错: {e}")
                results.append({
                    'content': '',
                    'format': 'unknown',
                    'metadata': {'error': str(e)},
                    'original_length': 0,
                    'cleaned_length': 0
                })
        return results


if __name__ == '__main__':
    print("开始流程")

    # 创建 ETL 服务实例
    etl_service = ETLService(use_unstructured=False)

    # 示例：处理文档（需要提供实际文件路径）
    # result = etl_service.process_document("path/to/document.pdf")
    # print(f"处理结果: {result['metadata']}")
    # print(f"清洗后内容长度: {len(result['content'])}")

    # 示例一：仅清洗文本
    # sample_text = """
    # <html>
    # <head><title>测试文档</title></head>
    # <body>
    #     <h1>这是标题</h1>
    #     <p>这是正文内容。包含一些特殊字符：@@@ ### $$$</p>
    #     <p>第 1 页</p>
    #     <p>2024-01-01</p>
    # </body>
    # </html>
    # """
    #
    # cleaned = etl_service.clean_text(sample_text)
    # print("原始文本:")
    # print(sample_text)
    # print("\n清洗后文本:")
    # print(cleaned)
    # 示例二：doc 文档清洗
    files_path = [os.path.dirname(__file__) + "\企业应急管理服务资料V1.0.0测试.docx"]
    print(files_path)
    parse_texts = etl_service.batch_process(files_path)
    for index, parse_text in enumerate(parse_texts):
        print(f"文档 {index}:")
        print(f"内容 {parse_text.get('content')}:")
